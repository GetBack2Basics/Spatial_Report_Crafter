#!/usr/bin/env python3
"""
create_extract_docx.py -- Generate Water_Sewer_Reconciliation_Extract.docx

Reads Critical_Review_Assets.gpkg (output of run_full_reconciliation_review.py)
and generates a structured Word document summarizing the reconciliation results
for Water, Sewer, and WaterMeters domains.

Usage:
    python create_extract_docx.py [--gpkg path/to/Critical_Review_Assets.gpkg]
"""

import argparse
import os
import sys
from pathlib import Path

import geopandas as gpd
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

WORKING_DIR = Path("/media/george-corea/GIS/Projects/014_CSC_Reflect/Working")
DEFAULT_GPKG = WORKING_DIR / "Pending Data" / "Critical_Review_Assets.gpkg"
OUT_DOCX = WORKING_DIR / "Data Catalogue" / "Water_Sewer_Reconciliation_Extract.docx"


# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

COLOR_PRIMARY = RGBColor(31, 78, 120)
COLOR_SECONDARY = RGBColor(46, 116, 181)
COLOR_TEXT = RGBColor(51, 51, 51)


def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def style_header_cell(cell, text, bg="1F4E78"):
    cell.text = text
    set_cell_background(cell, bg)
    set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Data loading
# ---------------------------------------------------------------------------

def load_review_data(gpkg_path):
    """Load all review layers from the GPKG and return as a dict of {layer_name: gdf}."""
    layers = {}
    import fiona
    for layer_name in fiona.listlayers(str(gpkg_path)):
        gdf = gpd.read_file(gpkg_path, layer=layer_name)
        layers[layer_name] = gdf
    return layers


def summarize_layers(layers):
    """Build summary statistics per domain."""
    domains = {
        "water_mains": {"new": None, "drift": None, "vectors": None},
        "water_meters": {"new": None, "drift": None, "vectors": None},
        "sewer_manholes": {"new": None, "drift": None, "vectors": None},
        "sewer_mains": {"new": None, "drift": None, "vectors": None},
    }

    for name, gdf in layers.items():
        if gdf.empty:
            continue
        name_lower = name.lower()
        if "water_meter" in name_lower:
            domain = "water_meters"
        elif "water_main" in name_lower:
            domain = "water_mains"
        elif "sewer_main" in name_lower:
            domain = "sewer_mains"
        elif "sewer_manhole" in name_lower:
            domain = "sewer_manholes"
        else:
            continue

        if "new_review" in name_lower:
            domains[domain]["new"] = gdf
        elif "drift_review" in name_lower:
            domains[domain]["drift"] = gdf
        elif "displacement_vector" in name_lower:
            domains[domain]["vectors"] = gdf

    return domains


# ---------------------------------------------------------------------------
# Document generation
# ---------------------------------------------------------------------------

def generate_document(domains, out_path):
    doc = Document()

    # Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Water & Sewer Asset Reconciliation: Full Review Results")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY

    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run("Reflect Project Asset Management Team\n")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = COLOR_SECONDARY

    # Executive Summary
    callout_table = doc.add_table(rows=1, cols=1)
    callout_table.style = "Table Grid"
    callout_cell = callout_table.cell(0, 0)
    set_cell_background(callout_cell, "F2F5F8")
    set_cell_margins(callout_cell, top=120, bottom=120, left=200, right=200)
    tcPr = callout_cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="36" w:space="0" w:color="1F4E78"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

    cp = callout_cell.paragraphs[0]
    r1 = cp.add_run("Executive Summary: ")
    r1.font.bold = True
    r1.font.color.rgb = COLOR_PRIMARY

    total_new = sum(len(d["new"]) if d["new"] is not None else 0 for d in domains.values())
    total_drift = sum(len(d["drift"]) if d["drift"] is not None else 0 for d in domains.values())
    r2 = cp.add_run(
        f"Full reconciliation of Reflect field data against GIS targets for Water, "
        f"Sewer, and WaterMeters domains. Found {total_new} new assets and "
        f"{total_drift} drift features across {len(domains)} asset classes."
    )
    r2.font.italic = True

    doc.add_paragraph()

    # Section 1: Summary Metrics Table
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Summary of Reconciliation Metrics")
    h1_run.font.color.rgb = COLOR_PRIMARY
    h1_run.font.size = Pt(15)

    table = doc.add_table(rows=5, cols=6)
    table.style = "Table Grid"

    headers = [
        "Asset Category", "Target GIS Layer", "New Assets",
        "Drift Features", "Displacement Vectors", "Total Flagged",
    ]
    for col_idx, text in enumerate(headers):
        style_header_cell(table.cell(0, col_idx), text)

    row_data = []
    for domain_key, data in domains.items():
        n_new = len(data["new"]) if data["new"] is not None else 0
        n_drift = len(data["drift"]) if data["drift"] is not None else 0
        n_vec = len(data["vectors"]) if data["vectors"] is not None else 0
        total = n_new + n_drift
        label = domain_key.replace("_", " ").title()
        row_data.append([label, domain_key, str(n_new), str(n_drift), str(n_vec), str(total)])

    for row_idx, rd in enumerate(row_data):
        for col_idx, text in enumerate(rd):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = text
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            if row_idx % 2 == 1:
                set_cell_background(cell, "F2F5F8")

    doc.add_paragraph()

    # Section 2: New Assets Detail
    h2 = doc.add_heading(level=1)
    h2_run = h2.add_run("2. New Assets (Missing from GIS)")
    h2_run.font.color.rgb = COLOR_PRIMARY
    h2_run.font.size = Pt(15)

    for domain_key, data in domains.items():
        new_gdf = data["new"]
        if new_gdf is None or new_gdf.empty:
            continue
        label = domain_key.replace("_", " ").title()
        doc.add_heading(f"{label} -- {len(new_gdf)} new assets", level=3)

        # Show first 50 in table, summarize rest
        display_gdf = new_gdf.head(50)
        cols_to_show = ["Drift_m", "Risk_Level", "Src_CSV", "Src_Type", "Src_Date", "Src_Value", "Comments"]
        available_cols = [c for c in cols_to_show if c in display_gdf.columns]

        if available_cols:
            t = doc.add_table(rows=len(display_gdf) + 1, cols=len(available_cols))
            t.style = "Table Grid"
            for ci, col_name in enumerate(available_cols):
                style_header_cell(t.cell(0, ci), col_name, bg="2E74B5")
            for ri, (_, row) in enumerate(display_gdf.iterrows()):
                for ci, col_name in enumerate(available_cols):
                    val = str(row.get(col_name, ""))
                    if len(val) > 80:
                        val = val[:77] + "..."
                    cell = t.cell(ri + 1, ci)
                    cell.text = val
                    if ri % 2 == 1:
                        set_cell_background(cell, "F2F5F8")

        if len(new_gdf) > 50:
            doc.add_paragraph(f"... and {len(new_gdf) - 50} more new assets (see GPKG for full list).")

        doc.add_paragraph()

    # Section 3: Drift Features Detail
    h3 = doc.add_heading(level=1)
    h3_run = h3.add_run("3. Drift Features (Spatial Deviations)")
    h3_run.font.color.rgb = COLOR_PRIMARY
    h3_run.font.size = Pt(15)

    for domain_key, data in domains.items():
        drift_gdf = data["drift"]
        if drift_gdf is None or drift_gdf.empty:
            continue
        label = domain_key.replace("_", " ").title()
        doc.add_heading(f"{label} -- {len(drift_gdf)} drift features", level=3)

        # Risk summary
        if "Risk_Level" in drift_gdf.columns:
            risk_counts = drift_gdf["Risk_Level"].value_counts()
            risk_p = doc.add_paragraph()
            risk_p.add_run("Risk Distribution: ").font.bold = True
            parts = [f"{v} {k}" for k, v in risk_counts.items()]
            risk_p.add_run("; ".join(parts))

        # Show high-risk items
        high_risk = drift_gdf[drift_gdf["Risk_Level"].str.contains("High", na=False)]
        if not high_risk.empty:
            doc.add_heading("High Risk Items (>10m drift)", level=4)
            cols_to_show = ["Drift_m", "Risk_Level", "Src_CSV", "Src_Type", "Review_PotentialChange", "Review_Status"]
            available_cols = [c for c in cols_to_show if c in high_risk.columns]

            if available_cols:
                t = doc.add_table(rows=min(len(high_risk), 50) + 1, cols=len(available_cols))
                t.style = "Table Grid"
                for ci, col_name in enumerate(available_cols):
                    style_header_cell(t.cell(0, ci), col_name, bg="C0392B")
                for ri, (_, row) in enumerate(high_risk.head(50).iterrows()):
                    for ci, col_name in enumerate(available_cols):
                        val = str(row.get(col_name, ""))
                        if len(val) > 80:
                            val = val[:77] + "..."
                        cell = t.cell(ri + 1, ci)
                        cell.text = val
                        if ri % 2 == 1:
                            set_cell_background(cell, "FDEDEC")

        doc.add_paragraph()

    # Section 4: Attribute Conflicts (Review_PotentialChange)
    h4 = doc.add_heading(level=1)
    h4_run = h4.add_run("4. Attribute Conflicts (Review_PotentialChange)")
    h4_run.font.color.rgb = COLOR_PRIMARY
    h4_run.font.size = Pt(15)

    has_conflicts = False
    for domain_key, data in domains.items():
        drift_gdf = data["drift"]
        if drift_gdf is None or drift_gdf.empty or "Review_PotentialChange" not in drift_gdf.columns:
            continue
        conflicts = drift_gdf[drift_gdf["Review_PotentialChange"].notna() & (drift_gdf["Review_PotentialChange"] != "")]
        if conflicts.empty:
            continue
        has_conflicts = True
        label = domain_key.replace("_", " ").title()
        doc.add_heading(f"{label} -- {len(conflicts)} conflicts", level=3)

        cols = ["Drift_m", "Risk_Level", "Review_PotentialChange", "Src_CSV"]
        available_cols = [c for c in cols if c in conflicts.columns]
        if available_cols:
            t = doc.add_table(rows=len(conflicts) + 1, cols=len(available_cols))
            t.style = "Table Grid"
            for ci, col_name in enumerate(available_cols):
                style_header_cell(t.cell(0, ci), col_name, bg="8E44AD")
            for ri, (_, row) in enumerate(conflicts.iterrows()):
                for ci, col_name in enumerate(available_cols):
                    val = str(row.get(col_name, ""))
                    if len(val) > 100:
                        val = val[:97] + "..."
                    cell = t.cell(ri + 1, ci)
                    cell.text = val
                    if ri % 2 == 1:
                        set_cell_background(cell, "F5EEF8")

        doc.add_paragraph()

    if not has_conflicts:
        doc.add_paragraph("No attribute conflicts detected between source and target.")

    # Section 5: Status Flags (Review_Status)
    h5 = doc.add_heading(level=1)
    h5_run = h5.add_run("5. Status Flags (Review_Status)")
    h5_run.font.color.rgb = COLOR_PRIMARY
    h5_run.font.size = Pt(15)

    has_status = False
    for domain_key, data in domains.items():
        for layer_type in ["new", "drift"]:
            gdf = data[layer_type]
            if gdf is None or gdf.empty or "Review_Status" not in gdf.columns:
                continue
            flagged = gdf[gdf["Review_Status"].notna() & (gdf["Review_Status"] != "")]
            if flagged.empty:
                continue
            has_status = True
            label = domain_key.replace("_", " ").title()
            doc.add_heading(f"{label} ({layer_type}) -- {len(flagged)} flagged", level=3)

            cols = ["Drift_m", "Risk_Level", "Review_Status", "Src_CSV"]
            available_cols = [c for c in cols if c in flagged.columns]
            if available_cols:
                t = doc.add_table(rows=len(flagged) + 1, cols=len(available_cols))
                t.style = "Table Grid"
                for ci, col_name in enumerate(available_cols):
                    style_header_cell(t.cell(0, ci), col_name, bg="E67E22")
                for ri, (_, row) in enumerate(flagged.iterrows()):
                    for ci, col_name in enumerate(available_cols):
                        val = str(row.get(col_name, ""))
                        if len(val) > 100:
                            val = val[:97] + "..."
                        cell = t.cell(ri + 1, ci)
                        cell.text = val
                        if ri % 2 == 1:
                            set_cell_background(cell, "FEF5E7")

            doc.add_paragraph()

    if not has_status:
        doc.add_paragraph("No status flags detected.")

    # Section 6: Next Steps
    h6 = doc.add_heading(level=1)
    h6_run = h6.add_run("6. Next Steps")
    h6_run.font.color.rgb = COLOR_PRIMARY
    h6_run.font.size = Pt(15)

    doc.add_paragraph("Based on this full review, the following actions are recommended:")

    steps = [
        ("QA Review in QGIS", "Open Critical_Review_Project.qgs to visually inspect all flagged features. "
         "Use the rule-based symbology to prioritize high-risk drift items (>10m, Crimson markers)."),
        ("Validate New Assets", "Cross-reference new asset IDs with field crew records to confirm "
         "they are genuine new installations and not duplicates or miscataloged fittings."),
        ("Resolve Attribute Conflicts", "Review items in Section 4 (Review_PotentialChange) to "
         "determine which source values should be promoted to the master GIS database."),
        ("Process Status Flags", "Review items in Section 5 (Review_Status) for assets marked "
         "Abandoned, Not Found, or Removed. Update GIS status fields accordingly."),
        ("Production Write-Back", "After QA approval, run the reconciliation script to auto-update "
         "GIS geometries (drift <= 10m) and append lineage columns."),
    ]

    for title, desc in steps:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{title}: ")
        r.font.bold = True
        p.add_run(desc)

    # Save
    doc.save(str(out_path))
    print(f"DOCX generated: {out_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Generate reconciliation DOCX report")
    parser.add_argument(
        "--gpkg", type=str, default=str(DEFAULT_GPKG),
        help="Path to Critical_Review_Assets.gpkg",
    )
    parser.add_argument(
        "--output", type=str, default=str(OUT_DOCX),
        help="Output DOCX path",
    )
    args = parser.parse_args()

    gpkg_path = Path(args.gpkg)
    if not gpkg_path.exists():
        print(f"ERROR: GPKG not found: {gpkg_path}")
        print("Run run_full_reconciliation_review.py first.")
        sys.exit(1)

    print(f"Loading review data from: {gpkg_path}")
    layers = load_review_data(gpkg_path)
    print(f"Found {len(layers)} layers")

    domains = summarize_layers(layers)
    for dk, dv in domains.items():
        n_new = len(dv["new"]) if dv["new"] is not None else 0
        n_drift = len(dv["drift"]) if dv["drift"] is not None else 0
        n_vec = len(dv["vectors"]) if dv["vectors"] is not None else 0
        print(f"  {dk}: {n_new} new, {n_drift} drift, {n_vec} vectors")

    generate_document(domains, args.output)


if __name__ == "__main__":
    main()
